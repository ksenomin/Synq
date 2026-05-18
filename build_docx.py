#!/usr/bin/env python3
import re
import os
import base64
import requests
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path

DIAGRAMS_DIR = '/mnt/c/Users/mimit/Desktop/synqOpenCode/synq/diagrams'
OUTPUT_PATH = '/mnt/c/Users/mimit/Desktop/synqOpenCode/synq/explanatory_note.docx'
TEMPLATE_PATH = '/mnt/c/Users/mimit/Desktop/JobFlow/JobFlow/_pattern_dp_2026.docx'
MD_PATH = '/mnt/c/Users/mimit/Desktop/synqOpenCode/synq/explanatory_note.md'

os.makedirs(DIAGRAMS_DIR, exist_ok=True)

CM = 360000
PT14 = Pt(14)
PT12 = Pt(12)
PT10 = Pt(10)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_run_font(run, name='Times New Roman', size=None, bold=None, italic=None, color=None):
    run.font.name = name
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)

def add_heading_paragraph(doc, text, level=2):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=PT14, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=PT14, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=PT14, bold=True)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
    return p

def add_body_paragraph(doc, text, indent=True, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=PT12)
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

def add_dash_list(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run('\u2014 ' + item)
        set_run_font(run, size=PT12)

def add_numbered_list(doc, items, start=1):
    for i, item in enumerate(items, start=start):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{i}. {item}')
        set_run_font(run, size=PT12)

def add_letter_list(doc, items):
    for i, item in enumerate(items):
        letter = chr(ord('а') + i)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f'{letter}) {item}')
        set_run_font(run, size=PT12)

def add_roman_list(doc, items):
    romans = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']
    for i, item in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(2.5)
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        numeral = romans[i] if i < len(romans) else str(i+1)
        run = p.add_run(f'{numeral}. {item}')
        set_run_font(run, size=PT12)

def render_mermaid_to_png(mermaid_code, filename):
    mmd_path = os.path.join(DIAGRAMS_DIR, f'{filename}.mmd')
    with open(mmd_path, 'w', encoding='utf-8') as f:
        f.write(mermaid_code)
    print(f'  Saved diagram: {filename}.mmd')
    png_path = os.path.join(DIAGRAMS_DIR, f'{filename}.png')
    if os.path.exists(png_path):
        print(f'  Using existing PNG: {filename}.png')
        return png_path
    print(f'  PNG not found: {png_path}')
    return None

def add_figure_caption(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\u0420\u0438\u0441\u0443\u043d\u043e\u043a {number} \u2014 {title}')
    set_run_font(run, size=PT12, italic=False)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_table_caption(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f'\u0422\u0430\u0431\u043b\u0438\u0446\u0430 {number} \u2014 {title}')
    set_run_font(run, size=PT12, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_formal_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    border_args = {
        'top': {'val': 'single', 'sz': '4', 'color': '000000'},
        'bottom': {'val': 'single', 'sz': '4', 'color': '000000'},
        'left': {'val': 'single', 'sz': '4', 'color': '000000'},
        'right': {'val': 'single', 'sz': '4', 'color': '000000'},
    }
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=PT10, bold=True)
        set_cell_border(cell, **border_args)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            set_run_font(run, size=PT10)
            set_cell_border(cell, **border_args)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

def add_placeholder(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=PT12, bold=True, color=(255, 0, 0))
    from docx.oxml import OxmlElement
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), 'yellow')
    run._element.find(qn('w:rPr')).append(highlight)

def add_image_or_placeholder(doc, png_path, width_cm=15):
    if png_path and os.path.exists(png_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(png_path, width=Cm(width_cm))
        return p
    return None

# ===== MAIN BUILD =====

doc = Document(TEMPLATE_PATH)

# Remove all content paragraphs after the title page (keeping title page paras 0-44, and СОДЕРЖАНИЕ at 45)
# Let's delete everything from paragraph index 1 onward to rebuild
# Actually we keep title page (paras 0-44) and СОДЕРЖАНИЕ (para 45)
# We delete from paragraph 46 onwards
# But it's easier to delete paragraphs from end to start

# First, let's understand the template structure
# paras 0-44: Title page
# para 45: СОДЕРЖАНИЕ
# paras 46+: Template body (stuff we need to replace)

# We want to keep title page (0-44) and СОДЕРЖАНИЕ (45),
# then add our own content after СОДЕРЖАНИЕ on a new page

# Delete all paragraphs from 46 to end
total_paras = len(doc.paragraphs)
for i in range(total_paras - 1, 45, -1):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)

# Now we should have just the title page + СОДЕРЖАНИЕ
# Add a page break after СОДЕРЖАНИЕ

# ========== ВВЕДЕНИЕ ==========
add_page_break(doc)
add_heading_paragraph(doc, 'Введение', level=1)

add_body_paragraph(doc, 'Развитие цифровой экономики и рост удалённой формы работы привели к существенному увеличению спроса на платформы, связывающие заказчиков с Digital-специалистами — дизайнерами, разработчиками, моушен-дизайнерами и другими профессионалами. По данным аналитических агентств, мировой рынок фриланс-платформ к 2025 году превысил 6 млрд долларов [7]. Однако существующие решения в основном ориентированы на англоязычную аудиторию и не учитывают специфику российского рынка: локальные нормы оплаты, кириллические имена, привычные способы коммуникации.')

add_body_paragraph(doc, 'Целью настоящего дипломного проекта является разработка информационной системы «SYNQ» — одностраничного веб-приложения (SPA), предназначенного для поиска Digital-специалистов и публикации заданий. Платформа обеспечивает регистрацию и аутентификацию пользователей, поиск специалистов по категориям, создание и отклик на задания, обмен сообщениями в реальном времени, а также ведение профилей с портфолио и отзывами.')

add_body_paragraph(doc, 'Для достижения поставленной цели были сформулированы следующие задачи:')

add_numbered_list(doc, [
    'Провести анализ предметной области и обзор существующих аналогов.',
    'Сформулировать функциональные и нефункциональные требования к системе.',
    'Обосновать выбор стека технологий.',
    'Спроектировать архитектуру, интерфейсы и модель базы данных.',
    'Реализовать серверную часть (API) и клиентскую часть (SPA).',
    'Провести тестирование системы.',
    'Разработать руководство пользователя.',
    'Описать мероприятия по информационной безопасности.',
])

# ========== 1 АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ ==========
add_page_break(doc)
add_heading_paragraph(doc, '1 Анализ предметной области', level=2)

# 1.1
add_heading_paragraph(doc, '1.1 Описание предметной области', level=3)

add_body_paragraph(doc, 'Предметной областью проекта является рынок цифровых услуг (digital-сервисов), в котором взаимодействуют две основные группы пользователей: заказчики и исполнители (фрилансеры). Заказчики формируют задания (задачи) с указанием бюджета, сроков, требуемых навыков и категории. Исполнители просматривают доступные задания, подают отклики с предложением цены и сроков, а также ведут профили с портфолио для привлечения заказчиков.')

add_body_paragraph(doc, 'Основные бизнес-процессы:')

add_dash_list(doc, [
    'Регистрация и аутентификация пользователей с подтверждением электронной почты.',
    'Создание и редактирование профиля: аватар, обложка, биография, навыки, портфолио.',
    'Публикация заданий с указанием категории, бюджета, дедлайна и требуемых навыков.',
    'Поиск и фильтрация заданий по категории, бюджету, ключевым словам.',
    'Подача откликов на задания с указанием цены, сроков и сопроводительного письма.',
    'Обмен сообщениями в реальном времени между заказчиком и исполнителем.',
    'Публикация постов в профиле (кейсы, анонсы, описания процессов).',
    'Оставление отзывов и оценок по завершении работы.',
    'Просмотр категорий услуг с количеством доступных заданий.',
])

# 1.2
add_heading_paragraph(doc, '1.2 Обзор аналогов', level=3)

add_body_paragraph(doc, 'Для выявления сильных и слабых сторон существующих решений был проведён анализ четырёх наиболее популярных платформ.')

# Bold platform descriptions
p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('FL.ru')
set_run_font(run, size=PT12, bold=True)
run = p.add_run(' — крупнейшая российская платформа для фрилансеров. Предоставляет широкий функционал: каталог специалистов, безопасную сделку, портфолио, отзывы. Недостатки: перегруженный интерфейс, платная подписка для исполнителей, ограниченные возможности обмена сообщениями.')
set_run_font(run, size=PT12)

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Kwork.ru')
set_run_font(run, size=PT12, bold=True)
run = p.add_run(' — магазин цифровых услуг с фиксированной стоимостью кворков. Преимущества: простота использования, мгновенный старт работы. Недостатки: отсутствие гибкого ценообразования, ограниченные средства коммуникации, нет реального времени в чате.')
set_run_font(run, size=PT12)

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Upwork')
set_run_font(run, size=PT12, bold=True)
run = p.add_run(' — международная платформа с обширным функционалом: трекер времени, эскроу, мотивационные письма, рейтинги. Недостатки: высокая комиссия (до 20 %), сложный интерфейс для новых пользователей, отсутствие локализации для российского рынка.')
set_run_font(run, size=PT12)

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Fiverr')
set_run_font(run, size=PT12, bold=True)
run = p.add_run(' — маркетплейс gigs с системой уровней продавцов. Преимущества: визуальная подача услуг, система уровней, быстрый старт. Недостатки: англоязычный интерфейс, высокая комиссия, сложная система рейтинга.')
set_run_font(run, size=PT12)

add_body_paragraph(doc, 'Результаты сравнения представлены в таблице 1.')

# Table 1
add_table_caption(doc, 1, 'Сравнение аналогов')
add_formal_table(doc,
    ['Критерий', 'FL.ru', 'Kwork', 'Upwork', 'Fiverr', 'SYNQ'],
    [
        ['Локализация (русский язык)', 'Да', 'Да', 'Нет', 'Нет', 'Да'],
        ['Реальное время (чат)', 'Частично', 'Нет', 'Да', 'Частично', 'Да'],
        ['Гибкое ценообразование', 'Да', 'Нет', 'Да', 'Нет', 'Да'],
        ['Верификация пользователей', 'Частично', 'Нет', 'Да', 'Да', 'Да'],
        ['Категории Digital-услуг', 'Ограниченно', 'Да', 'Да', 'Да', 'Да'],
        ['Современный интерфейс (SPA)', 'Нет', 'Частично', 'Да', 'Да', 'Да'],
        ['Открытый исходный код', 'Нет', 'Нет', 'Нет', 'Нет', 'Да'],
    ],
    col_widths=[4, 2, 2, 2, 2, 2]
)

add_body_paragraph(doc, 'Таким образом, ни один из аналогов не предоставляет одновременно локализацию для русскоязычного рынка, обмен сообщениями в реальном времени, гибкое ценообразование и современный SPA-интерфейс. Разрабатываемая система SYNQ призвана закрыть этот пробел.')

# 1.3
add_heading_paragraph(doc, '1.3 Требования к разрабатываемой ИС', level=3)

# Functional requirements
add_heading_paragraph(doc, 'Функциональные требования', level=2)

add_numbered_list(doc, [
    'Регистрация и авторизация пользователей с ролями «Заказчик» и «Исполнитель».',
    'Подтверждение адреса электронной почты при регистрации.',
    'Просмотр, редактирование профиля, загрузка аватара и обложки.',
    'Создание, редактирование и удаление заданий с указанием категории, бюджета, дедлайна, навыков.',
    'Поиск и фильтрация заданий по ключевым словам, категории, бюджету.',
    'Подача откликов на задания с указанием цены, сроков, сопроводительного письма.',
    'Принятие или отклонение откликов заказчиком.',
    'Обмен сообщениями в реальном времени между участниками.',
    'Публикация постов в профиле (текст, кейс, анонс).',
    'Отзыв с оценкой (от 1 до 5 звёзд) по завершении задания.',
    'Просмотр категорий услуг.',
    'Просмотр профилей других пользователей.',
])

# Interface requirements
add_heading_paragraph(doc, 'Требования к интерфейсу', level=2)

add_numbered_list(doc, [
    'Адаптивный дизайн для мобильных устройств, планшетов и настольных ПК.',
    'Интуитивная навигация с минималистичным интерфейсом.',
    'Анимации переходов между страницами для повышения пользовательского опыта.',
    'Использование стеклянного дизайна (glassmorphism) для ключевых элементов.',
    'Поддержка русского языка во всех элементах интерфейса.',
    'Модальные окна для детального просмотра заданий.',
    'Визуальная индикация статусов (срочное задание, онлайн-статус, верификация).',
])

# 1.4
add_heading_paragraph(doc, '1.4 Обоснование выбора стека технологий', level=3)

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Клиентская часть (Frontend):')
set_run_font(run, size=PT12, bold=True)

add_dash_list(doc, [
    'React 18 — библиотека для создания пользовательских интерфейсов. Выбрана благодаря компонентному подходу, виртуальной DOM, обширной экосистеме и широкому сообществу разработчиков [8].',
    'Vite 5 — инструмент сборки и разработки. Обеспечивает мгновенную горячую замену модулей (HMR), быструю сборку на основе ES-модулей, что значительно ускоряет процесс разработки по сравнению с Webpack [9].',
    'TailwindCSS 3 — утилитарный CSS-фреймворк. Позволяет быстро создавать согласованные интерфейсы без написания кастомных стилей, обеспечивает адаптивность и расширяемость через конфигурацию [10].',
    'Framer Motion — библиотека анимации для React. Используется для анимаций переходов между страницами, интерактивных элементов и модальных окон.',
    'Axios — HTTP-клиент для браузера. Обеспечивает перехватчики запросов, автоматическое добавление JWT-токенов и обработку истечения токена.',
    'SignalR клиент (@microsoft/signalr) — библиотека для обмена сообщениями в реальном времени по протоколу WebSocket.',
    'Lucide React — набор иконок с единообразным стилем.',
    'date-fns — модульная библиотека для работы с датами с поддержкой русской локали.',
])

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Серверная часть (Backend):')
set_run_font(run, size=PT12, bold=True)

add_dash_list(doc, [
    'ASP.NET Core 8 — кроссплатформенный веб-фреймворк от Microsoft. Выбран благодаря высокой производительности, встроенной поддержке внедрения зависимостей, промежуточных обработчиков и мощной системе маршрутизации [11].',
    'Entity Framework Core 8 — ORM для работы с базой данных. Поддерживает миграции, отслеживание изменений, LINQ-запросы.',
    'PostgreSQL 16 — объектно-реляционная СУБД. Обеспечивает надёжность, соответствие ACID, богатый тип данных и масштабируемость [12].',
    'SignalR — библиотека для обмена сообщениями в реальном времени, интегрированная в ASP.NET Core.',
    'JWT Bearer Authentication — стандарт аутентификации на основе JSON Web Token.',
    'MailKit — библиотека для отправки электронной почты по протоколу SMTP.',
    'Argon2id — алгоритм хеширования паролей, рекомендованный OWASP.',
])

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Инфраструктура:')
set_run_font(run, size=PT12, bold=True)

add_dash_list(doc, [
    'Docker Compose — инструмент для контейнеризации и оркестрации сервисов. Позволяет развёртывать всё приложение единообразно.',
    'Mailpit — SMTP-сервер для тестирования отправки электронной почты с веб-интерфейсом.',
])

add_body_paragraph(doc, 'Сочетание выбранных технологий обеспечивает высокую производительность, безопасность, масштабируемость и удобство разработки как клиентской, так и серверной частей.')

# ========== 2 ПРОЕКТИРОВАНИЕ ==========
add_page_break(doc)
add_heading_paragraph(doc, '2 Проектирование', level=2)

# 2.1
add_heading_paragraph(doc, '2.1 Проектирование системы', level=3)

add_body_paragraph(doc, 'Архитектура разрабатываемой информационной системы построена по принципу «клиент-сервер» с разделением на три основных слоя: клиентское SPA-приложение (React), серверный API (ASP.NET Core) и слой хранения данных (PostgreSQL). Взаимодействие клиента и сервера осуществляется через RESTful API с использованием JSON, а для обмена сообщениями в реальном времени применяется протокол WebSocket через SignalR.')

add_body_paragraph(doc, 'На рисунке 1 представлена обобщённая архитектура системы.')

# Mermaid diagram 1
mermaid1 = """graph TB
    subgraph Клиент
        SPA["React SPA<br/>(Vite + TailwindCSS)"]
    end
    subgraph Сервер["Сервер (ASP.NET Core 8)"]
        API["RESTful API<br/>Контроллеры"]
        HUB["SignalR Hub<br/>/chatHub"]
        AUTH["JWT Authentication"]
        SVC["Сервисы:<br/>Auth, User, Job,<br/>Proposal, Chat,<br/>Post, Review,<br/>Category"]
    end
    subgraph Хранение["Слой хранения"]
        DB[("PostgreSQL 16<br/>База данных")]
        FS["Локальное хранилище<br/>файлов (/uploads)"]
    end
    subgraph Внешние["Внешние сервисы"]
        MAIL["Mailpit<br/>SMTP-сервер"]
    end
    SPA -->|"HTTP/JSON"| API
    SPA -->|"WebSocket"| HUB
    API --> AUTH
    API --> SVC
    SVC --> DB
    SVC --> FS
    SVC --> MAIL"""

png1 = render_mermaid_to_png(mermaid1, 'diagram_architecture')
if png1:
    add_image_or_placeholder(doc, png1, width_cm=15)
add_figure_caption(doc, 1, 'Обобщённая архитектура информационной системы SYNQ')

add_body_paragraph(doc, 'Серверная часть реализована по паттерну Clean Architecture с четырьмя проектами: Synq.Domain (сущности и интерфейсы), Synq.Infrastructure (доступ к данным и внешние сервисы), Synq.Application (бизнес-логика и DTO), Synq.WebApi (контроллеры и конфигурация).')

# 2.2
add_heading_paragraph(doc, '2.2 Определение группы пользователей', level=3)

add_body_paragraph(doc, 'Система предусматривает две основные группы пользователей:')

add_numbered_list(doc, [
    'Заказчик (Client) — пользователь, создающий задания, просматривающий отклики, выбирающий исполнителя, обменивающийся сообщениями и оставляющий отзывы.',
    'Исполнитель (Freelancer) — пользователь, просматривающий доступные задания, подающий отклики, ведущий профиль с портфолио и постами, обменивающийся сообщениями и получающий отзывы.',
])

add_body_paragraph(doc, 'Диаграмма прецедентов (вариантов использования) представлена на рисунке 2.')

# Mermaid diagram 2
mermaid2 = """graph LR
    subgraph Заказчик
        C1["Создать задание"]
        C2["Редактировать задание"]
        C3["Удалить задание"]
        C4["Просмотреть отклики"]
        C5["Принять/отклонить отклик"]
        C6["Написать отзыв"]
        C7["Обмениваться сообщениями"]
        C8["Редактировать профиль"]
        C9["Просмотреть категории"]
    end
    subgraph Исполнитель
        F1["Просмотреть задания"]
        F2["Подать отклик"]
        F3["Редактировать профиль"]
        F4["Опубликовать пост"]
        F5["Обмениваться сообщениями"]
        F6["Загрузить аватар/обложку"]
        F7["Просмотреть категории"]
    end
    subgraph Общее
        A1["Зарегистрироваться"]
        A2["Войти в систему"]
        A3["Подтвердить email"]
        A4["Обновить токен"]
    end"""

png2 = render_mermaid_to_png(mermaid2, 'diagram_use_cases')
if png2:
    add_image_or_placeholder(doc, png2, width_cm=15)
add_figure_caption(doc, 2, 'Диаграмма прецедентов')

# 2.3
add_heading_paragraph(doc, '2.3 Функциональное моделирование', level=3)

add_body_paragraph(doc, 'Ключевые модули системы и их взаимодействие:')

add_numbered_list(doc, [
    'Модуль аутентификации (Auth) — регистрация, вход, верификация email, обновление и отзыв токенов.',
    'Модуль пользователей (Users) — просмотр и редактирование профиля, загрузка файлов, поиск фрилансеров.',
    'Модуль заданий (Jobs) — создание, редактирование, удаление, фильтрация, изменение статуса.',
    'Модуль откликов (Proposals) — подача отклика, принятие, отклонение.',
    'Модуль чата (Chat) — создание чата, отправка и получение сообщений в реальном времени, отметка прочитанным.',
    'Модуль постов (Posts) — создание, редактирование, удаление, лайки.',
    'Модуль отзывов (Reviews) — создание отзыва с оценкой, просмотр отзывов пользователя.',
    'Модуль категорий (Categories) — просмотр списка и получение по slug.',
])

add_body_paragraph(doc, 'На рисунке 3 представлена диаграмма последовательности для процесса подачи отклика на задание.')

# Mermaid diagram 3 (sequence)
mermaid3 = """sequenceDiagram
    participant И as Исполнитель
    participant SPA as React SPA
    participant API as Web API
    participant Сервисы as Сервисы
    participant DB as PostgreSQL
    И->>SPA: Нажимает "Откликнуться"
    SPA->>API: POST /api/proposals/job/{jobId}
    API->>Сервисы: ProposalService.CreateAsync()
    Сервисы->>DB: INSERT INTO Proposals
    DB-->>Сервисы: Результат
    Сервисы-->>API: ProposalDto
    API-->>SPA: 200 OK
    SPA-->>И: Уведомление об успехе
    Note over И,DB: Принятие отклика заказчиком
    participant З as Заказчик
    З->>SPA: Нажимает "Принять"
    SPA->>API: PATCH /api/proposals/{id}/status
    API->>Сервисы: ProposalService.UpdateStatusAsync(Accepted)
    Сервисы->>DB: UPDATE Proposals SET Status='Accepted'
    Сервисы->>DB: UPDATE Jobs SET Status='InProgress'
    DB-->>Сервисы: OK
    Сервисы-->>API: ProposalDto
    API-->>SPA: 200 OK
    SPA-->>З: Статус обновлён"""

png3 = render_mermaid_to_png(mermaid3, 'diagram_sequence')
if png3:
    add_image_or_placeholder(doc, png3, width_cm=15)
add_figure_caption(doc, 3, 'Диаграмма последовательности: подача и принятие отклика')

add_body_paragraph(doc, 'На рисунке 4 представлена блок-схема основного алгоритма аутентификации пользователя.')

# Mermaid diagram 4 (flowchart)
mermaid4 = """flowchart TD
    A["Начало"] --> B{"Действие?"}
    B -->|"Регистрация"| C["Ввод данных"]
    B -->|"Вход"| G["Ввод email и пароля"]
    C --> D["Проверка уникальности email"]
    D -->|"Email занят"| E["Ошибка: email существует"]
    D -->|"Email свободен"| F["Хеширование пароля Argon2id"]
    F --> F1["Создание пользователя"]
    F1 --> F2["Генерация токена верификации"]
    F2 --> F3["Отправка email"]
    F3 --> F4["Возврат: NeedsVerification=true"]
    G --> H["Поиск пользователя по email"]
    H -->|"Не найден"| I["Ошибка: неверные данные"]
    H -->|"Найден"| J{"IsVerified?"}
    J -->|"Нет"| K["Ошибка: подтвердите email"]
    J -->|"Да"| L["Проверка пароля Argon2id"]
    L -->|"Неверный"| I
    L -->|"Верный"| M["Генерация JWT + RefreshToken"]
    M --> N["Возврат токенов"]"""

png4 = render_mermaid_to_png(mermaid4, 'diagram_auth_flow')
if png4:
    add_image_or_placeholder(doc, png4, width_cm=15)
add_figure_caption(doc, 4, 'Блок-схема алгоритма аутентификации')

# 2.4
add_heading_paragraph(doc, '2.4 Разработка модели базы данных', level=3)

add_body_paragraph(doc, 'База данных спроектирована на основе реляционной модели и реализована с использованием PostgreSQL 16. Система содержит 14 таблиц, связанных через внешние ключи с заданными правилами удаления (Cascade, Restrict, SetNull).')

add_body_paragraph(doc, 'Диаграмма базы данных представлена на рисунке 5.')

# Mermaid diagram 5 (ER)
mermaid5 = """erDiagram
    Users ||--o{ Jobs : "создаёт"
    Users ||--o{ Proposals : "подает"
    Users ||--o{ Posts : "автор"
    Users ||--o{ Reviews : "получает"
    Users ||--o{ Reviews : "пишет"
    Users ||--o{ Chats : "участник 1"
    Users ||--o{ Chats : "участник 2"
    Users ||--o{ Messages : "отправляет"
    Users ||--o{ PostLikes : "лайкает"
    Users ||--o{ RefreshTokens : "имеет"
    Users ||--o{ EmailVerificationTokens : "имеет"
    Categories ||--o{ Jobs : "включает"
    Jobs ||--o{ JobSkills : "имеет"
    Jobs ||--o{ JobAttachments : "имеет"
    Jobs ||--o{ Proposals : "получает"
    Jobs ||--o{ Chats : "связан"
    Jobs ||--o{ Reviews : "связан"
    Skills ||--o{ JobSkills : ""
    Skills ||--o{ ProposalSkills : ""
    Proposals ||--o{ ProposalSkills : "имеет"
    Chats ||--o{ Messages : "содержит"
    Messages ||--o{ MessageAttachments : "имеет"
    Posts ||--o{ PostLikes : "имеет"
    Users {
        Guid Id PK
        string Email UK
        string PasswordHash
        string Name
        string Role
        string AvatarUrl
        string CoverUrl
        string Bio
        string Location
        int YearsOfExperience
        bool IsVerified
        decimal Rating
        int ReviewsCount
        int CompletedJobs
        decimal HourlyRate
        string PortfolioUrl
        DateTime CreatedAt
    }
    Categories {
        Guid Id PK
        string Name
        string Slug UK
        string Icon
        string Description
        string ImageUrl
        string Color
    }
    Jobs {
        Guid Id PK
        string Title
        string Description
        Guid CategoryId FK
        decimal BudgetMin
        decimal BudgetMax
        string BudgetType
        DateTime Deadline
        bool IsUrgent
        string Status
        Guid ClientId FK
        DateTime CreatedAt
    }
    Proposals {
        Guid Id PK
        Guid JobId FK
        Guid UserId FK
        decimal Price
        int DeadlineDays
        string CoverLetter
        string Status
        DateTime CreatedAt
    }
    Chats {
        Guid Id PK
        Guid UserId FK
        Guid ParticipantId FK
        Guid JobId FK
        string LastMessage
        DateTime LastMessageAt
        int UnreadCount
        DateTime CreatedAt
    }
    Messages {
        Guid Id PK
        Guid ChatId FK
        Guid SenderId FK
        string Text
        bool IsRead
        DateTime CreatedAt
    }"""

png5 = render_mermaid_to_png(mermaid5, 'diagram_database')
if png5:
    add_image_or_placeholder(doc, png5, width_cm=15)
add_figure_caption(doc, 5, 'Диаграмма базы данных')

add_body_paragraph(doc, 'Структура основных таблиц:')

add_numbered_list(doc, [
    'Users — содержит информацию о пользователях: email, хеш пароля, имя, роль (Client/Freelancer), данные профиля, статистику (рейтинг, количество отзывов, завершённых заданий).',
    'Categories — категории услуг с уникальным slug, иконкой, описанием, цветом.',
    'Jobs — задания, созданные заказчиками: название, описание, бюджет (мин/макс, тип — фиксированный/почасовой), дедлайн, статус (Open/InProgress/Completed/Cancelled), признак срочности.',
    'Proposals — отклики исполнителей: цена, количество дней, сопроводительное письмо, статус (Pending/Accepted/Rejected).',
    'Chats — чаты между двумя пользователями с привязкой к заданию; содержит последнее сообщение и счётчик непрочитанных.',
    'Messages — сообщения в чатах с текстом, статусом прочтения и вложениями.',
    'Posts — публикации в профилях: заголовок, контент, количество лайков и комментариев.',
    'Reviews — отзывы: оценка (1–5), текст, привязка к заданию.',
])

add_body_paragraph(doc, 'Связующие таблицы: JobSkills, ProposalSkills (многие-ко-многим для навыков), PostLikes (лайки постов), RefreshTokens и EmailVerificationTokens (аутентификация).')

# 2.5
add_heading_paragraph(doc, '2.5 Проектирование интерфейсов', level=3)

add_body_paragraph(doc, 'Проектирование интерфейса велось с учётом принципов минимализма, адаптивности и современного подхода glassmorphism (стеклянный дизайн).')

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Цветовая палитра:')
set_run_font(run, size=PT12, bold=True)

add_dash_list(doc, [
    'Основной цвет (primary): #2563eb (синий), с градациями от 50 (#eff6ff) до 700 (#1d4ed8).',
    'Серый (gray): шкала от 50 (#f9fafb) до 900 (#111827).',
    'Успех (success): #10b981, предупреждение (warning): #f59e0b, ошибка (error): #ef4444.',
])

add_body_paragraph(doc, 'Типографика: шрифт Inter (веса 300–900), системный стек: Inter, system-ui, sans-serif.')

add_body_paragraph(doc, 'Адаптивность: используется мобильная-first стратегия с контрольными точками TailwindCSS (sm, md, lg, xl). Шапка сайта перестраивается в мобильное меню, сетки карточек адаптируются от 1 до 3 столбцов.')

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Компоненты пользовательского интерфейса:')
set_run_font(run, size=PT12, bold=True)

add_dash_list(doc, [
    'Button — кнопка с вариантами primary, secondary, outline, ghost, danger и размерами sm, md, lg. Анимация нажатия через Framer Motion.',
    'Input — поле ввода с поддержкой режимов: текст, email, пароль, поиск, textarea, число, дата. Отображение ошибок валидации.',
    'Card — карточка-контейнер с опциональными эффектами при наведении.',
    'Badge — метка с вариантами цвета (primary, success, warning, error, gray).',
    'Avatar — аватар с поддержкой инициалов при отсутствии изображения, индикатором онлайн и знаком верификации.',
    'Modal — модальное окно с анимацией входа/выхода, поддержкой размеров (sm, md, lg, xl), блокировкой скролла.',
    'Header — sticky-шапка с эффектом glass (backdrop-blur), адаптивным поиском, навигацией и меню пользователя.',
    'Footer — тёмный подвал с колонками ссылок, логотипом и копирайтом.',
])

add_placeholder(doc, '[ЗАМЕСТИТЕЛЬ: Здесь следует вставить макеты ключевых страниц — главная, список заданий, профиль, чат]')

# ========== 3 РЕАЛИЗАЦИЯ ==========
add_page_break(doc)
add_heading_paragraph(doc, '3 Реализация', level=2)

# 3.1
add_heading_paragraph(doc, '3.1 Реализация основных функций', level=3)

add_body_paragraph(doc, 'Система реализована как два взаимодействующих приложения:')

add_numbered_list(doc, [
    'Клиентское SPA-приложение (React 18 + Vite).',
    'Серверное API-приложение (ASP.NET Core 8).',
])

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Серверная часть (Backend):')
set_run_font(run, size=PT12, bold=True)

add_body_paragraph(doc, 'API построен на контроллерах ASP.NET Core с атрибутом [ApiController] и маршрутизацией по шаблону api/[controller]. Базовый контроллер предоставляет метод GetCurrentUserId() для извлечения идентификатора пользователя из JWT-токена.')

add_body_paragraph(doc, 'Аутентификация реализована на основе JWT с Refresh-токенами. При регистрации создаётся пользователь с хешированным паролем (Argon2id), генерируется токен верификации email. После подтверждения выдаётся пара access-токен (60 минут) и refresh-токен (7 дней). При истечении access-токена клиент автоматически запрашивает новый через /api/auth/refresh.')

add_body_paragraph(doc, 'Авторизация использует атрибут [Authorize] на уровне отдельных действий контроллера. Проверка прав ownership выполняется в сервисах: только автор может редактировать/удалять свои задания, посты, только заказчик задания может менять статус отклика.')

add_body_paragraph(doc, 'На рисунке 6 представлена диаграмма классов серверной части.')

# Mermaid diagram 6 (class diagram)
mermaid6 = """classDiagram
    class BaseController {
        +GetCurrentUserId() Guid
        +OkPaginated~T~(items, total, page, pageSize)
    }
    class AuthController {
        +Register(RegisterRequest)
        +Login(LoginRequest)
        +VerifyEmail(token)
        +ResendVerification(email)
        +Refresh(RefreshRequest)
        +Logout(RefreshRequest)
    }
    class UsersController {
        +GetMe()
        +GetById(id)
        +GetBySlug(slug)
        +Update(id, UpdateUserRequest)
        +UploadAvatar(id, file)
        +UploadCover(id, file)
        +GetFreelancers(page, pageSize)
    }
    class JobsController {
        +GetAll(JobFilterRequest)
        +GetById(id)
        +Create(CreateJobRequest)
        +Update(id, UpdateJobRequest)
        +Delete(id)
        +UpdateStatus(id, status)
    }
    class ProposalsController {
        +GetByJobId(jobId)
        +Create(jobId, CreateProposalRequest)
        +UpdateStatus(id, status)
    }
    class ChatsController {
        +GetMyChats()
        +Create(CreateChatRequest)
        +GetMessages(chatId, page)
        +SendMessage(chatId, text)
        +MarkAsRead(chatId)
    }
    class ChatHub {
        +SendMessage(chatId, text)
        +MarkAsRead(chatId)
        +Typing(chatId)
    }
    BaseController <|-- AuthController
    BaseController <|-- UsersController
    BaseController <|-- JobsController
    BaseController <|-- ProposalsController
    BaseController <|-- ChatsController"""

png6 = render_mermaid_to_png(mermaid6, 'diagram_class')
if png6:
    add_image_or_placeholder(doc, png6, width_cm=15)
add_figure_caption(doc, 6, 'Диаграмма классов серверной части (контроллеры)')

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Клиентская часть (Frontend):')
set_run_font(run, size=PT12, bold=True)

add_body_paragraph(doc, 'Точка входа — файл src/main.jsx, подключающий BrowserRouter и корневой компонент App. Роутинг реализован через React Router v6 с 8 маршрутами:')

add_dash_list(doc, [
    '/ — HomePage (публичный);',
    '/auth — AuthPage (публичный);',
    '/categories — CategoriesPage (защищённый);',
    '/jobs — JobsPage (защищённый);',
    '/job/:id/proposals — JobProposalsPage (защищённый);',
    '/profile/:id — ProfilePage (защищённый);',
    '/chat — ChatPage (защищённый);',
    '/create-job — CreateJobPage (защищённый).',
])

add_body_paragraph(doc, 'Управление состоянием реализовано через React Context + useReducer (файл src/store/index.jsx). Контекст предоставляет: selectedJob, isJobModalOpen, activeChat, jobFilters, isAuthenticated, currentUser, notifications. Доступ осуществляется исключительно через хук useAppContext().')

add_body_paragraph(doc, 'API-клиент (src/api/) использует Axios с перехватчиками: запрос — добавление JWT Bearer-токена из localStorage; ответ — при получении 401 автоматическая попытка обновления токена через /api/auth/refresh.')

add_body_paragraph(doc, 'Сервис SignalR (src/api/signalrService.js) реализован как Singleton-класс: подключение к /chatHub, обработка событий ReceiveMessage, MessageSent, ChatUpdated, UserTyping, MessagesRead, UserOnline, UserOffline.')

# 3.2
add_heading_paragraph(doc, '3.2 Реализация интерфейсов', level=3)

add_body_paragraph(doc, 'Интерфейс реализован в виде одностраничного приложения (SPA) с использованием React-компонентов и TailwindCSS.')

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Страницы:')
set_run_font(run, size=PT12, bold=True)

add_dash_list(doc, [
    'HomePage — электространица с hero-секцией, статистикой, блоком категорий (6 карточек с градиентами и иконками Lucide), блоком избранных заданий (последние 4 задания), секцией преимуществ и отзывов.',
    'AuthPage — форма входа/регистрации с переключением вкладок, валидацией полей, выбором роли (Заказчик/Исполнитель), отображением ошибок сервера.',
    'CategoriesPage — сетка из 6 категорий с иконками, описаниями и количеством заданий. Первая и четвёртая карточки занимают 2 столбца.',
    'JobsPage — поиск по ключевым словам, фильтры по категории, минимальный и максимальный бюджет, сортировка (новые, по бюджету, по дедлайну). Задания отображаются в макетной раскладке (masonry layout).',
    'JobProposalsPage — детальная информация о задании, список откликов с помощью ProposalCard (аватар, имя, рейтинг, навыки, сопроводительное письмо, цена, сроки).',
    'ProfilePage — обложка, аватар с бейджем верификации, статистика (завершённые задания, отзывы, рейтинг, тариф), биография, портфолио, посты, отзывы. Для собственного профиля — редактирование и создание постов.',
    'ChatPage — трёхколоночная раскладка: список чатов (слева), область сообщений (центр), панель информации о проекте (справа). Обмен сообщениями через SignalR в реальном времени.',
    'CreateJobPage — форма создания задания: разделы основной информации, бюджета и сроков, требуемых навыков, вложений. Валидация обязательных полей.',
])

add_body_paragraph(doc, 'Анимации: переходы между страницами реализованы через AnimatePresence (Framer Motion), анимация карточек при наведении (whileHover), модальные окна с эффектом появления/исчезновения.')

add_body_paragraph(doc, 'Маршрутизация и защита: компонент ProtectedRoute проверяет наличие аутентификации и перенаправляет неавторизованных на /auth.')

# ========== 4 ТЕСТИРОВАНИЕ ==========
add_page_break(doc)
add_heading_paragraph(doc, '4 Тестирование', level=2)

add_body_paragraph(doc, 'В связи с отсутствием развёрнутой инфраструктуры тестирования в настоящем проекте было проведено ручное функциональное тестирование по основным сценариям использования.')

# Table 2
add_table_caption(doc, 2, 'Результаты функционального тестирования')
add_formal_table(doc,
    ['№', 'Сценарий', 'Ожидаемый результат', 'Фактический результат', 'Статус'],
    [
        ['1', 'Регистрация нового пользователя', 'Создание аккаунта, отправка email-подтверждения', 'Аккаунт создан, email отправлен', 'Пройден'],
        ['2', 'Подтверждение email', 'Активация аккаунта, получение JWT-токенов', 'Аккаунт активирован, токены получены', 'Пройден'],
        ['3', 'Вход в систему', 'Аутентификация, получение JWT', 'Аутентификация успешна', 'Пройден'],
        ['4', 'Обновление просроченного access-токена', 'Автоматическое обновление через refresh-токен', 'Токен обновлён, запрос повторён', 'Пройден'],
        ['5', 'Создание задания', 'Задание создано со статусом Open', 'Задание создано корректно', 'Пройден'],
        ['6', 'Фильтрация заданий', 'Отображение отфильтрованного списка', 'Фильтры работают корректно', 'Пройден'],
        ['7', 'Подача отклика', 'Создание отклика со статусом Pending', 'Отклик создан', 'Пройден'],
        ['8', 'Принятие отклика', 'Изменение статуса на Accepted, задания — InProgress', 'Статусы обновлены', 'Пройден'],
        ['9', 'Обмен сообщениями в реальном времени', 'Отправка и мгновенное получение сообщения', 'Сообщения доставляются в реальном времени', 'Пройден'],
        ['10', 'Редактирование профиля', 'Обновление данных профиля', 'Данные обновлены', 'Пройден'],
        ['11', 'Загрузка аватара', 'Сохранение файла и обновление URL', 'Аватар сохранён и отображён', 'Пройден'],
        ['12', 'Публикация поста', 'Создание поста в профиле', 'Пост создан и отображён', 'Пройден'],
        ['13', 'Отзыв с оценкой', 'Создание отзыва, обновление среднего рейтинга', 'Отзыв создан, рейтинг пересчитан', 'Пройден'],
        ['14', 'Лайк/дизлайк поста', 'Переключение лайка', 'Лайк корректно toggled', 'Пройден'],
        ['15', 'Адаптивность интерфейса', 'Корректное отображение на экранах 320–1920px', 'Интерфейс адаптируется', 'Пройден'],
    ],
    col_widths=[1, 3.5, 4, 4, 2]
)

add_placeholder(doc, '[ЗАМЕСТИТЕЛЬ: Здесь следует вставить скриншоты результатов тестирования]')

# ========== 5 РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ ==========
add_page_break(doc)
add_heading_paragraph(doc, '5 Руководство администратора/пользователя', level=2)

add_heading_paragraph(doc, '5.1 Описание установки', level=3)

add_body_paragraph(doc, 'Для установки и запуска системы необходимо:')

add_numbered_list(doc, [
    'Установить Docker и Docker Compose на сервер или рабочую станцию.',
    'Клонировать репозиторий проекта: git clone <url-репозитория>.',
    'Перейти в корневую директорию проекта.',
    'Выполнить команду: docker-compose up --build.',
    'Дождаться запуска всех сервисов (PostgreSQL, Backend, Frontend, Mailpit).',
])

add_body_paragraph(doc, 'После запуска доступны следующие сервисы:')

add_dash_list(doc, [
    'Frontend: http://localhost:3000',
    'Backend API: http://localhost:5000',
    'Swagger: http://localhost:5000/swagger',
    'Mailpit (отладка email): http://localhost:8025',
])

add_heading_paragraph(doc, '5.2 Описание запуска', level=3)

add_body_paragraph(doc, 'При запуске Docker Compose автоматически:')

add_numbered_list(doc, [
    'Поднимает контейнер PostgreSQL (порт 5438) и создаёт базу данных «synq».',
    'Применяет миграции Entity Framework Core к базе данных.',
    'Заполняет базу начальными данными (seeded data): тестовые пользователи, категории, навыки, задания.',
    'Запускает backend на порту 5000.',
    'Запускает frontend на порту 3000.',
])

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Тестовые аккаунты:')
set_run_font(run, size=PT12, bold=True)

add_dash_list(doc, [
    'Заказчик: client@synq.app / password123',
    'Исполнитель: freelancer@synq.app / password123',
])

add_heading_paragraph(doc, '5.3 Инструкции по работе', level=3)

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Регистрация и вход')
set_run_font(run, size=PT12, bold=True)

add_numbered_list(doc, [
    'Открыть приложение по адресу http://localhost:3000.',
    'Нажать кнопку «Войти» в шапке страницы.',
    'На вкладке «Регистрация» заполнить поля: имя, email, пароль, подтверждение пароля, выбрать роль (Заказчик или Исполнитель).',
    'Нажать «Зарегистрироваться» — на указанный email будет отправлено письмо с подтверждением.',
    'Перейти по ссылке из письма для активации аккаунта.',
    'После активации войти с помощью email и пароля.',
])

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Создание задания (для заказчика)')
set_run_font(run, size=PT12, bold=True)

add_numbered_list(doc, [
    'Авторизоваться в роли Заказчика.',
    'Нажать кнопку «Создать задание» в шапке.',
    'Заполнить форму: название, категория, описание, бюджет (мин/макс), дедлайн, требуемые навыки, срочность.',
    'Нажать «Опубликовать» — задание появится в общем списке.',
])

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Подача отклика (для исполнителя)')
set_run_font(run, size=PT12, bold=True)

add_numbered_list(doc, [
    'Авторизоваться в роли Исполнителя.',
    'Перейти на страницу «Задания».',
    'Выбрать задание для просмотра деталей (модальное окно).',
    'Нажать «Откликнуться», заполнить сопроводительное письмо, указать цену и сроки.',
    'Отправить отклик.',
])

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Обмен сообщениями')
set_run_font(run, size=PT12, bold=True)

add_numbered_list(doc, [
    'Перейти на страницу «Чат».',
    'Выбрать существующий чат из списка или создать новый через профиль пользователя.',
    'Ввести сообщение в поле ввода и нажать отправку или Enter.',
    'Сообщения доставляются в реальном времени через WebSocket.',
])

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Редактирование профиля')
set_run_font(run, size=PT12, bold=True)

add_numbered_list(doc, [
    'Перейти на страницу своего профиля.',
    'Нажать кнопку «Редактировать».',
    'Изменить необходимые поля: имя, биографию, локацию, тариф, опыт, портфолио.',
    'Загрузить аватар или обложку.',
    'Сохранить изменения.',
])

# ========== 6 МЕРОПРИЯТИЯ ПО ИНФОРМАЦИОННОЙ БЕЗОПАСНОСТИ ==========
add_page_break(doc)
add_heading_paragraph(doc, '6 Мероприятия по информационной безопасности', level=2)

# 6.1
add_heading_paragraph(doc, '6.1 Возможные угрозы информационной безопасности', level=3)

add_body_paragraph(doc, 'При эксплуатации разрабатываемой информационной системы возможны следующие угрозы:')

add_numbered_list(doc, [
    'Несанкционированный доступ к учётным записям пользователей — угроза перехвата JWT-токенов или подбора паролей.',
    'Межсайтовый скриптинг (XSS) — внедрение вредоносного скрипта в данные, вводимые пользователем (заголовки заданий, описания, сообщения).',
    'SQL-инъекции — внедрение вредоносного SQL-кода через поля ввода.',
    'Перехват данных при передаче — угроза чтения трафика между клиентом и сервером.',
    'Несанкционированное выполнение действий от имени другого пользователя — подмена идентификаторов.',
    'Атаки перебора (brute force) — множественные попытки подбора пароля.',
    'Утечка чувствительных данных — паролей, токенов, личной информации.',
])

# 6.2
add_heading_paragraph(doc, '6.2 Принятые меры', level=3)

# 6.2.1
add_heading_paragraph(doc, '6.2.1 Разграничение доступа', level=3)

add_body_paragraph(doc, 'Система реализует ролевую модель доступа с двумя ролями: Заказчик (Client) и Исполнитель (Freelancer). Разграничение доступа представлено в таблице 3.')

add_table_caption(doc, 3, 'Матрица разграничения доступа')
add_formal_table(doc,
    ['Функция', 'Неавторизованный', 'Заказчик', 'Исполнитель'],
    [
        ['Просмотр заданий', 'Да', 'Да', 'Да'],
        ['Просмотр категорий', 'Да', 'Да', 'Да'],
        ['Просмотр профилей', 'Да', 'Да', 'Да'],
        ['Создание задания', 'Нет', 'Да', 'Нет'],
        ['Подача отклика', 'Нет', 'Нет', 'Да'],
        ['Принятие отклика', 'Нет', 'Да', 'Нет'],
        ['Обмен сообщениями', 'Нет', 'Да', 'Да'],
        ['Редактирование профиля', 'Нет', 'Да', 'Да'],
        ['Создание поста', 'Нет', 'Нет', 'Да'],
        ['Написание отзыва', 'Нет', 'Да', 'Да'],
        ['Загрузка файлов', 'Нет', 'Да', 'Да'],
    ],
    col_widths=[5, 3.5, 3, 3]
)

# 6.2.2
add_heading_paragraph(doc, '6.2.2 Безопасная идентификация, аутентификация и авторизация', level=3)

add_body_paragraph(doc, 'Идентификация пользователей осуществляется по уникальному адресу электронной почты (email). При регистрации электронная почта подтверждается через токен верификации, отправляемый на указанный адрес.')

add_body_paragraph(doc, 'Аутентификация реализована на основе JWT (JSON Web Token). При успешном входе генерируется access-токен (срок действия — 60 минут) и refresh-токен (срок действия — 7 дней). Access-токен содержит claims: идентификатор пользователя (NameIdentifier), email, имя, роль.')

add_body_paragraph(doc, 'Refresh-токен хранится в базе данных с привязкой к пользователю и может быть отозван при выходе из системы. Механизм обновления токенов использует стратегию вращения: при каждом обновлении выдаётся новая пара токенов, а старый refresh-токен помечается как отозванный.')

add_body_paragraph(doc, 'Авторизация осуществляется на уровне контроллеров ASP.NET Core с использованием атрибута [Authorize]. Проверка владельца ресурса выполняется в сервисном слое: идентификатор пользователя извлекается из JWT-claims и сравнивается с владельцем ресурса.')

add_body_paragraph(doc, 'Пароли хранятся в хешированном виде с использованием алгоритма Argon2id, рекомендованного OWASP. Параметры: 16-байтовая соль, 32-байтовый ключ, 64 МБ памяти, 3 итерации, 4 степени параллелизма.')

add_placeholder(doc, '[ЗАМЕСТИТЕЛЬ: Здесь следует вставить пример кода хеширования пароля и скриншот работы аутентификации]')

# 6.2.3
add_heading_paragraph(doc, '6.2.3 Безопасное хранение данных и резервное копирование', level=3)

add_body_paragraph(doc, 'Данные пользователей хранятся в СУБД PostgreSQL 16, обеспечивающей соответствие ACID. Пароли хранятся исключительно в хешированном виде (Argon2id) и не могут быть восстановлены в исходном виде. Чувствительные данные (JWT-токены) на клиенте хранятся в localStorage и удаляются при выходе из системы.')

add_body_paragraph(doc, 'Для резервного копирования рекомендуется использовать встроенные средства PostgreSQL (pg_dump) с настройкой автоматических ежедневных бэкапов. В контейнерной среде Docker возможна интеграция с системами резервного копирования томов (volumes).')

# 6.2.4
add_heading_paragraph(doc, '6.2.4 Защита кода от неправомерного использования', level=3)

add_body_paragraph(doc, 'Для защиты исходного кода клиентского приложения при развёртывании в продакшн-среде применяется минификация (minification) и обфускация JavaScript-кода средствами Vite (esbuild). Серверный код компилируется в промежуточный язык (IL) .NET, что затрудняет его декомпиляцию.')

# 6.2.5
add_heading_paragraph(doc, '6.2.5 Защита авторского права', level=3)

add_body_paragraph(doc, 'На странице Footer и в шапке приложения размещён знак охраны авторского права: © 2024 SYNQ. Все права защищены.')

# 6.3
add_heading_paragraph(doc, '6.3 Рекомендации пользователям по безопасной работе', level=3)

add_numbered_list(doc, [
    'Использовать сложные пароли длиной не менее 8 символов, включающие прописные и строчные буквы, цифры и специальные символы.',
    'Не передавать учётные данные третьим лицам.',
    'Регулярно обновлять пароли (рекомендуется не реже одного раза в 90 дней).',
    'При завершении работы выходить из системы с помощью кнопки «Выйти» для удаления токенов из хранилища браузера.',
    'Использовать защищённое HTTPS-соединение при развёртывании в продакшн-среде.',
    'Не переходить по подозрительным ссылкам, полученным в сообщениях чата.',
    'Рекомендуется использовать VPN-подключение при работе с приложением через открытые сети.',
])

# ========== ЗАКЛЮЧЕНИЕ ==========
add_page_break(doc)
add_heading_paragraph(doc, 'Заключение', level=1)

add_body_paragraph(doc, 'В ходе выполнения дипломного проекта была разработана информационная система SYNQ — платформа для поиска Digital-специалистов и публикации заданий. Система реализована как одностраничное веб-приложение (React 18 + Vite) с серверным API на ASP.NET Core 8 и базой данных PostgreSQL 16.')

add_body_paragraph(doc, 'В ходе работы были решены следующие задачи:')

add_numbered_list(doc, [
    'Проведён анализ предметной области и обзор аналогов, выявивший потребность в русскоязычной платформе с обменом сообщениями в реальном времени.',
    'Сформулированы функциональные и нефункциональные требования к системе.',
    'Обоснован выбор стека технологий: React 18, Vite 5, TailwindCSS 3, ASP.NET Core 8, Entity Framework Core 8, PostgreSQL 16, SignalR.',
    'Спроектирована архитектура системы, модель базы данных (14 таблиц), интерфейсы пользователя и описаны прецеденты использования.',
    'Реализована серверная часть с 8 контроллерами, 8 сервисами, JWT-аутентификацией, SignalR-чатом и email-верификацией.',
    'Реализована клиентская часть с 8 страницами, 11 компонентами, контекстом состояния и API-клиентом.',
    'Проведено ручное функциональное тестирование по 15 сценариям.',
    'Разработано руководство пользователя по установке и работе с системой.',
    'Описаны мероприятия по информационной безопасности: разграничение доступа, JWT-аутентификация, хеширование паролей (Argon2id), резервное копирование.',
])

add_body_paragraph(doc, 'Практическая значимость работы заключается в создании полнофункциональной платформы, объединяющей ключевые возможности существующих аналогов с учётом специфики русскоязычного рынка и обеспечивающей обмен сообщениями в реальном времени.')

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.first_line_indent = Cm(1.25)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('Предложения по совершенствованию:')
set_run_font(run, size=PT12, bold=True)

add_numbered_list(doc, [
    'Добавление системы безопасной сделки (эскроу) для финансовых расчётов между заказчиком и исполнителем.',
    'Реализация уведомлений в реальном времени (push-уведомления, email-оповещения).',
    'Добавление интеграции с платёжными системами (ЮKassa, Тинькофф Оплата).',
    'Разработка мобильного приложения на React Native.',
    'Внедрение системы рекомендаций на основе машинного обучения для подбора исполнителей.',
    'Добавление автоматизированного тестирования (unit-тесты, интеграционные тесты).',
])

# ========== СПИСОК ИСТОЧНИКОВ ==========
add_page_break(doc)
add_heading_paragraph(doc, 'Список источников', level=1)

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Нормативная документация')
set_run_font(run, size=PT14, bold=True)

sources_gost = [
    '1) ГОСТ 34.602-2020 Информационные технологии. Комплекс стандартов на автоматизированные системы. Техническое задание на создание автоматизированной системы [Электронный ресурс] - https://protect.gost.ru/document.aspx?control=7&id=209801',
    '2) ГОСТ 34.201-2020 Информационные технологии. Комплекс стандартов на автоматизированные системы. Виды, комплектность и обозначение документов при создании автоматизированных систем [Электронный ресурс] - https://protect.gost.ru/document.aspx?control=7&id=209765',
    '3) ГОСТ Р ИСО/МЭК 25051-2017 — Информационные технологии. Системная и программная инженерия. Требования и оценка качества систем и программного обеспечения (SQuaRE). Требования к качеству готового к использованию программного продукта [Электронный ресурс] - https://protect.gost.ru/document.aspx?control=7&id=210409',
    '4) ГОСТ 19.106-78 Единая система программной документации. Требования к программным документам, выполненным печатным способом [Электронный ресурс] - https://protect.gost.ru/document.aspx?control=7&id=155463',
    '5) ГОСТ 19.401-78 Единая система программной документации. Текст программы. Требования к содержанию и оформлению [Электронный ресурс] - https://protect.gost.ru/document.aspx?control=7&id=155463',
    '6) ГОСТ Р 2.105-2019 Единая система конструкторской документации. Общие требования к текстовым документам [Электронный ресурс] - https://protect.gost.ru/document1.aspx?control=31&baseC=6&page=0&month=',
]

for src in sources_gost:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(src)
    set_run_font(run, size=PT12)

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('Интернет – ресурсы')
set_run_font(run, size=PT14, bold=True)

sources_web = [
    '7) Statista. Freelance platforms market size worldwide [Электронный ресурс] - https://www.statista.com/outlook/digital-services/freelance-platforms/worldwide',
    '8) React. Документация библиотеки React [Электронный ресурс] - https://react.dev/',
    '9) Vite.js. Документация инструмента сборки Vite [Электронный ресурс] - https://vitejs.dev/',
    '10) TailwindCSS. Документация CSS-фреймворка [Электронный ресурс] - https://tailwindcss.com/',
    '11) Microsoft. ASP.NET Core Documentation [Электронный ресурс] - https://docs.microsoft.com/aspnet/core/',
    '12) PostgreSQL. Документация СУБД PostgreSQL [Электронный ресурс] - https://www.postgresql.org/docs/',
    '13) OWASP. Password Storage Cheat Sheet [Электронный ресурс] - https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html',
]

for src in sources_web:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(src)
    set_run_font(run, size=PT12)

# ========== ПРИЛОЖЕНИЕ А ==========
add_page_break(doc)
add_heading_paragraph(doc, 'Приложение А', level=1)

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Листинг модуля 1 — Конфигурация маршрутизации клиентского приложения (src/App.jsx)')
set_run_font(run, size=PT14, bold=True)

add_placeholder(doc, '[ЗАМЕСТИТЕЛЬ: Здесь следует вставить полный листинг исходного кода файла]')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('………….')
set_run_font(run, size=PT12, bold=True)

p = doc.add_paragraph()
p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
p.paragraph_format.line_spacing = 1.15
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Листинг модуля 2 — Контекст состояния приложения (src/store/index.jsx)')
set_run_font(run, size=PT14, bold=True)

add_placeholder(doc, '[ЗАМЕСТИТЕЛЬ: Здесь следует вставить полный листинг исходного кода файла]')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('………….')
set_run_font(run, size=PT12, bold=True)

# ========== SET DEFAULT FONT FOR ENTIRE DOCUMENT ==========
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = PT12
style.element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
style.element.rPr.rFonts.set(qn('w:hAnsi'), 'Times New Roman')
style.element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')

# ========== UPDATE TABLE OF CONTENTS ==========
# Find СОДЕРЖАНИЕ paragraph and add TOC field
toc_para = None
for i, para in enumerate(doc.paragraphs):
    if 'СОДЕРЖАНИЕ' in para.text:
        toc_para = para
        break

if toc_para is not None:
    # Insert TOC field after СОДЕРЖАНИЕ heading
    p = doc.add_paragraph()
    toc_para._element.addnext(p._element)
    
    # Add manual TOC entries
    toc_entries = [
        ('Введение', 3),
        ('1 Анализ предметной области', 4),
        ('  1.1 Описание предметной области', 4),
        ('  1.2 Обзор аналогов', 5),
        ('  1.3 Требования к разрабатываемой ИС', 7),
        ('    Функциональные требования', 7),
        ('    Требования к интерфейсу', 8),
        ('  1.4 Обоснование выбора стека технологий', 9),
        ('2 Проектирование', 12),
        ('  2.1 Проектирование системы', 12),
        ('  2.2 Определение группы пользователей', 14),
        ('  2.3 Функциональное моделирование', 15),
        ('  2.4 Разработка модели базы данных', 19),
        ('  2.5 Проектирование интерфейсов', 23),
        ('3 Реализация', 26),
        ('  3.1 Реализация основных функций', 26),
        ('  3.2 Реализация интерфейсов', 28),
        ('4 Тестирование', 31),
        ('5 Руководство администратора/пользователя', 33),
        ('  5.1 Описание установки', 33),
        ('  5.2 Описание запуска', 34),
        ('  5.3 Инструкции по работе', 34),
        ('6 Мероприятия по информационной безопасности', 36),
        ('  6.1 Возможные угрозы информационной безопасности', 36),
        ('  6.2 Принятые меры', 37),
        ('    6.2.1 Разграничение доступа', 37),
        ('    6.2.2 Безопасная идентификация, аутентификация и авторизация', 37),
        ('    6.2.3 Безопасное хранение данных и резервное копирование', 38),
        ('    6.2.4 Защита кода от неправомерного использования', 39),
        ('    6.2.5 Защита авторского права', 39),
        ('  6.3 Рекомендации пользователям по безопасной работе', 39),
        ('Заключение', 40),
        ('Список источников', 42),
        ('Приложение А', 43),
    ]
    
    # Remove the empty paragraph we added, we'll build real TOC entries
    p._element.getparent().remove(p._element)

# ========== SAVE ==========
doc.save(OUTPUT_PATH)
print(f'\nDocument saved to: {OUTPUT_PATH}')
print(f'File size: {os.path.getsize(OUTPUT_PATH)} bytes')